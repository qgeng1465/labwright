# -*- coding: utf-8 -*-
"""paper/md_to_docx.py — faithful markdown -> submission .docx renderer.

Turns ``paper/manuscript.md`` into the submission writeup
``paper/Labwright_bench_copilot_writeup.docx``.

Faithfulness contract (project honesty protocol):

* Every string that reaches the docx is copied verbatim from ``manuscript.md``;
  the renderer adds *only* layout — fonts, margins, page numbers, table grid,
  and the four referenced figures embedded at their caption positions. No
  number, caption or sentence is generated here.
* Figure embedding: a paragraph of the form ``**Figure N.** (in
  `paper/fig_X.png`, rendered by ...) — caption`` has the parenthetical
  annotation stripped and the PNG embedded above the caption text.
* Inline math ``$...$`` is transliterated to Unicode runs (italic variables,
  sub/superscripts); constructs the translator does not recognise are left
  verbatim in the run so nothing is silently dropped.

Regenerate after editing manuscript.md::

    python -m paper.md_to_docx

The .docx is gitignored (submission-in-progress); this script is tracked so
the submission is reproducible from the markdown source.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "paper" / "manuscript.md"
OUT = ROOT / "paper" / "Labwright_bench_copilot_writeup.docx"

BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"

BLACK = RGBColor(0, 0, 0)

# ---------------------------------------------------------------------------
# small helpers
# ---------------------------------------------------------------------------


def _set_run_font(run, name=BODY_FONT, size=11, bold=False, italic=False,
                  sub=False, sup=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.subscript = sub
    run.font.superscript = sup
    run.font.color.rgb = color


def _shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _left_border(p, color="999999", sz="4"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _base_para(doc, *, align=None, before=0, after=6, indent_in=None,
               hanging_in=None, justify=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent_in:
        pf.left_indent = Inches(indent_in)
    if hanging_in:
        pf.left_indent = Inches(hanging_in)
        pf.first_line_indent = Inches(-hanging_in)
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


# ---------------------------------------------------------------------------
# inline markdown -> runs
# ---------------------------------------------------------------------------

_GREEK = {
    "\\alpha": "\u03b1", "\\beta": "\u03b2", "\\gamma": "\u03b3",
    "\\delta": "\u03b4", "\\epsilon": "\u03b5", "\\mu": "\u03bc",
    "\\nu": "\u03bd", "\\tau": "\u03c4", "\\pi": "\u03c0",
    "\\rho": "\u03c1", "\\sigma": "\u03c3", "\\phi": "\u03c6",
    "\\theta": "\u03b8", "\\omega": "\u03c9", "\\Delta": "\u0394",
    "\\Omega": "\u03a9", "\\times": "\u00d7", "\\cdot": "\u00b7",
    "\\pm": "\u00b1", "\\lambda": "\u03bb",
}


def _math_segments(latex: str):
    """Turn a ``$...$`` body into (text, style) segments.

    Only the constructs actually used in manuscript.md are translated; anything
    else is emitted verbatim so no math is silently dropped.
    """
    s = latex.strip()
    s = re.sub(r"\\frac\{([^{}]*)\}\{([^{}]*)\}", r"(\1)/(\2)", s)
    for k, v in _GREEK.items():
        s = s.replace(k, v)
    s = s.replace("\\left", "").replace("\\right", "")
    # LaTeX token-separator spaces (e.g. '\mu Q', 'w h^2', '\Delta P') collapse
    # in typeset output; remove them so the docx reads like the equation, not
    # the source markup.
    s = re.sub(r"\s+", "", s)

    segs: list[tuple[str, dict]] = []
    buf = ""
    i = 0
    while i < len(s):
        ch = s[i]
        if ch in "_^" and i + 1 < len(s):
            style = {"sub": ch == "_", "sup": ch == "^"}
            j = i + 1
            if s[j] == "{":
                k = s.find("}", j + 1)
                content = s[j + 1:k]
                i = k + 1
            else:
                content = s[j]
                i = j + 1
            if buf:
                segs.append((buf, {"sub": False, "sup": False}))
                buf = ""
            segs.append((content, style))
        else:
            buf += ch
            i += 1
    if buf:
        segs.append((buf, {"sub": False, "sup": False}))
    return segs


_INLINE = re.compile(
    r"(?P<star>\\\*)"                      # 1 literal asterisk (footnote marker)
    r"|(?P<bold>\*\*(?P<b>.+?)\*\*)"       # 2 bold
    # 3 italic — content may contain an escaped '\*' (footnote marker inside
    #   the span, e.g. the affiliation line); a bare '*' always closes.
    r"|(?P<ital>\*(?P<i>(?:\\\*|[^*])*?)\*)"
    r"|(?P<code>`(?P<c>[^`]+)`)"           # 4 inline code
    r"|(?P<math>\$(?P<m>[^$]+)\$)"         # 5 inline math
    r"|(?P<link>\[(?P<lt>[^\]]+)\]\((?P<lu>[^)]+)\))"  # 6 markdown link
)


def render_inline(p, text: str, size: int = 11, bold: bool = False,
                  italic: bool = False):
    """Add runs to paragraph ``p`` for inline markdown in ``text``."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            _set_run_font(r, size=size, bold=bold, italic=italic)
        kind = m.lastgroup
        if kind == "star":
            r = p.add_run("*")
            _set_run_font(r, size=size, bold=bold, italic=italic)
        elif kind == "bold":
            render_inline(p, m.group("b"), size=size, bold=True, italic=italic)
        elif kind == "ital":
            # recurse so inner `code` / $math$ / **bold** inside an italic
            # span (e.g. a table caption) render too, not as literal markers
            render_inline(p, m.group("i"), size=size, bold=bold, italic=True)
        elif kind == "code":
            r = p.add_run(m.group("c"))
            _set_run_font(r, name=CODE_FONT, size=max(size - 1, 8),
                          bold=bold, italic=italic)
        elif kind == "math":
            for seg_text, style in _math_segments(m.group("m")):
                r = p.add_run(seg_text)
                _set_run_font(r, size=size, bold=bold, italic=True,
                              sub=style["sub"], sup=style["sup"])
        elif kind == "link":
            r = p.add_run(m.group("lt"))
            _set_run_font(r, size=size, bold=bold, italic=italic)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        _set_run_font(r, size=size, bold=bold, italic=italic)


# ---------------------------------------------------------------------------
# block renderers
# ---------------------------------------------------------------------------

def _add_heading(doc, text, level: int):
    p = _base_para(doc, before=12 if level == 1 else 10, after=6)
    r = p.add_run(text)
    _set_run_font(r, size=14 if level == 1 else 12, bold=True)
    # keep the outline level so a navigable TOC can be built in Word
    pPr = p._p.get_or_add_pPr()
    out = OxmlElement("w:outlineLvl")
    out.set(qn("w:val"), str(level - 1))
    pPr.append(out)
    return p


def _render_blockquote(doc, qlines):
    content = " ".join(
        ln.lstrip("> ").strip() for ln in qlines if ln.strip() != ">"
    )
    # Abstract special-case: '> **Abstract.** ...' -> 'Abstract' heading + body.
    m = re.match(r"^\*\*Abstract\.\*\*(.*)$", content, re.S)
    if m:
        _add_heading(doc, "Abstract", 1)
        body = m.group(1).strip()
        if body:
            p = _base_para(doc, justify=True)
            render_inline(p, body)
        return
    # General blockquote: indented, thin left rule.
    for ln in [content] if content else []:
        p = _base_para(doc, indent_in=0.4, after=8, justify=True)
        _left_border(p)
        render_inline(p, ln)
        if not ln:
            return


def _render_table(doc, rows):
    # rows: list of '|...|' lines; drop the separator row (|---|).
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    # Drop the markdown separator row (|---|). A separator cell is all
    # hyphens/dashes with optional leading/trailing ':' (alignment marker).
    def _is_separator(cells):
        return bool(cells) and all(
            re.fullmatch(r":?-+:?", c or "") for c in cells
        )

    if len(parsed) >= 2 and _is_separator(parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.autofit = True
    for r_i, cells in enumerate(parsed):
        row_cells = table.add_row().cells
        for c_i in range(ncols):
            txt = cells[c_i] if c_i < len(cells) else ""
            cell = row_cells[c_i]
            cell.paragraphs[0].text = ""
            render_inline(cell.paragraphs[0], txt, size=9, bold=(r_i == 0))
            if r_i == 0:
                _shade(cell.paragraphs[0], "EFEFEF")
    doc.add_paragraph()


_FIG = re.compile(r"\(in\s+`([^`]+\.png)`")


def _render_paragraph_block(doc, lines):
    if not lines:
        return
    joined = " ".join(ln.strip() for ln in lines if ln.strip())
    # Figure caption with embedded PNG: '**Figure N.** (in `paper/fig_X.png`…)'.
    if re.match(r"^\*\*Figure\s+\d+\.\*\*", joined) and _FIG.search(joined):
        png = _FIG.search(joined).group(1)
        img_path = ROOT / png
        if img_path.exists():
            p_img = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
            with Image.open(img_path) as im:
                w_px, h_px = im.size
            target_w, target_h = 6.3, 7.2
            scale = min(target_w / (w_px / 96.0), target_h / (h_px / 96.0))
            p_img.add_run().add_picture(str(img_path), width=Inches((w_px / 96.0) * scale))
        caption = re.sub(r"\(\s*in\s+`[^`]+`[^)]*\)", "", joined)
        caption = re.sub(r"\s{2,}", " ", caption).strip()
        p = _base_para(doc, after=10, justify=True)
        render_inline(p, caption)
        return
    # Bullet / numbered lists (source already carries the marker text).
    bullet = re.match(r"^-\s", joined)
    number = re.match(r"^(\d+)\.\s", joined)
    if bullet or number:
        # Accumulate the whole item first, then render once: inline spans
        # (**bold**, *italic*, `code`, $math$) may break across source lines,
        # and rendering line-by-line would leave the opening marker on one line
        # and the closing marker on the next, i.e. literal '*'.
        cur_text = None
        for ln in lines:
            s = ln.strip()
            if not s:
                continue
            if re.match(r"^-\s", s) or re.match(r"^\d+\.\s", s):
                if cur_text is not None:
                    p = _base_para(doc, hanging_in=0.28, after=4, justify=True)
                    render_inline(p, cur_text)
                cur_text = s
            elif cur_text is not None:
                cur_text += " " + s
        if cur_text is not None:
            p = _base_para(doc, hanging_in=0.28, after=4, justify=True)
            render_inline(p, cur_text)
        return
    # Plain paragraph.
    p = _base_para(doc, justify=True)
    render_inline(p, joined)


def _render_code_block(doc, code_lines):
    for ln in code_lines:
        p = _base_para(doc, before=0, after=0)
        _shade(p)
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(ln if ln else " ")
        _set_run_font(r, name=CODE_FONT, size=9)
    _base_para(doc, after=6)


# ---------------------------------------------------------------------------
# document assembly
# ---------------------------------------------------------------------------

def _frontmatter(text: str) -> dict:
    fm = {}
    m = re.match(r"^---\s*\n(.*?)\n---", text, re.S)
    if not m:
        return fm
    for line in m.group(1).splitlines():
        if ":" in line:
            k, v = line.split(":", 1)
            fm[k.strip()] = v.strip().strip('"').strip("'")
    return fm


def main() -> int:
    src_text = SRC.read_text(encoding="utf-8")
    fm = _frontmatter(src_text)

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(1.0)
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")

    # footer page number
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    frun._r.append(f1); frun._r.append(it); frun._r.append(f2)

    # ---- title block ----
    body_text = re.sub(r"^---\s*\n.*?\n---\s*\n", "", src_text, count=1, flags=re.S)
    lines = body_text.splitlines()

    title = fm.get("title", "").strip()
    if title:
        p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        r = p.add_run(title)
        _set_run_font(r, size=15, bold=True)
    # Author comes from the body's own author line ('Q. Geng*') — the
    # frontmatter 'author' value is inconsistent ('Q., Geng').
    m_author = re.search(r"\*\*(?P<a>[^*\n]+?)\*\*\\\*", body_text)
    author = m_author.group("a").strip() if m_author else ""
    if author:
        p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        render_inline(p, "**" + author + "**\\*")
    aff = fm.get("affiliation", "")
    if aff:
        p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        render_inline(p, "*" + aff + ". \\*Correspondence: qgeng1465@users.noreply.github.com*")
    date = fm.get("date", "")
    if date:
        p = _base_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        r = p.add_run(date)
        _set_run_font(r, size=11)

    # ---- body ----
    i = 0
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()
        if not s:
            i += 1
            continue
        if s.startswith("```"):
            j = i + 1
            code = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code.append(lines[j])
                j += 1
            _render_code_block(doc, code)
            i = j + 1
        elif s.startswith(">"):
            j = i
            qlines = []
            while j < len(lines) and lines[j].strip().startswith(">"):
                qlines.append(lines[j])
                j += 1
            _render_blockquote(doc, qlines)
            i = j
        elif s.startswith("|"):
            j = i
            rows = []
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(lines[j].strip())
                j += 1
            _render_table(doc, rows)
            i = j
        elif s.startswith("## "):
            _add_heading(doc, s[3:].strip(), 2)
            i += 1
        elif s.startswith("### "):
            _add_heading(doc, s[4:].strip(), 3)
            i += 1
        elif s.startswith("# "):
            # the top-level title is already rendered on the title block
            i += 1
        else:
            j = i
            block = []
            while j < len(lines):
                s2 = lines[j].strip()
                if not s2:
                    break
                if s2.startswith(("```", ">", "|", "#")):
                    break
                block.append(lines[j])
                j += 1
            first = block[0].strip() if block else ""
            if (first == "**Q. Geng**\\*"
                    or first.startswith("*Department of Biomedical Engineering")):
                # already rendered on the title block
                i = j
                continue
            _render_paragraph_block(doc, block)
            i = j

    doc.save(str(OUT))
    print(f"saved -> {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
