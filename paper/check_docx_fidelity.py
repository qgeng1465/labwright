# -*- coding: utf-8 -*-
"""paper/check_docx_fidelity.py — verify the .docx faithfully reproduces manuscript.md.

Mirrors md_to_docx.py's block-splitting and inline normalization, then asserts
every block's normalized text appears verbatim in the regenerated docx text
(paragraphs + table cells). Any MISS is a rendering drift — a sentence, number
or caption that changed or vanished between source and submission docx.

Usage:
    python -m paper.check_docx_fidelity
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

from paper.md_to_docx import _INLINE, _math_segments, SRC, OUT

_FIG_ANNOT = re.compile(r"\(\s*in\s+`[^`]+`[^)]*\)")


def _norm_inline(text: str) -> str:
    """Text-equivalent of what render_inline() writes into runs."""
    parts = []
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            parts.append(text[pos:m.start()])
        kind = m.lastgroup
        if kind == "star":
            parts.append("*")
        elif kind == "bold":
            parts.append(_norm_inline(m.group("b")))
        elif kind == "ital":
            parts.append(_norm_inline(m.group("i")))
        elif kind == "code":
            parts.append(m.group("c"))
        elif kind == "math":
            parts.append("".join(t for t, _ in _math_segments(m.group("m"))))
        elif kind == "link":
            parts.append(m.group("lt"))
        pos = m.end()
    if pos < len(text):
        parts.append(text[pos:])
    return "".join(parts)


def _docx_text() -> str:
    doc = Document(str(OUT))
    chunks = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def _body_blocks(text: str):
    """Split manuscript.md (frontmatter stripped) into blocks exactly like
    md_to_docx.main(): each element is (kind, expected_substring_list)."""
    body = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, count=1, flags=re.S)
    lines = body.splitlines()
    blocks = []  # (kind, [expected text pieces])
    i = 0
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()
        if not s:
            i += 1
            continue
        if s.startswith("```"):
            j = i + 1
            code = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code.append(lines[j])
                j += 1
            blocks.append(("code", [ln for ln in code if ln.strip()]))
            i = j + 1
        elif s.startswith(">"):
            j = i
            qlines = []
            while j < len(lines) and lines[j].strip().startswith(">"):
                qlines.append(lines[j])
                j += 1
            content = " ".join(
                ln.lstrip("> ").strip() for ln in qlines if ln.strip() != ">"
            )
            m = re.match(r"^\*\*Abstract\.\*\*(.*)$", content, re.S)
            if m:
                body_txt = m.group(1).strip()
                pieces = ["Abstract"]
                if body_txt:
                    pieces.append(_norm_inline(body_txt))
            else:
                pieces = [_norm_inline(content)] if content else []
            blocks.append(("blockquote", pieces))
            i = j
        elif s.startswith("|"):
            j = i
            rows = []
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(lines[j].strip())
                j += 1
            parsed = []
            for row in rows:
                parsed.append([c.strip() for c in row.strip().strip("|").split("|")])
            if len(parsed) >= 2 and all(
                re.fullmatch(r":?-+:?", c or "") for c in parsed[1]
            ):
                parsed.pop(1)
            cells = []
            for r_i, row in enumerate(parsed):
                for c_i in range(len(row)):
                    cells.append(_norm_inline(row[c_i]))
            blocks.append(("table", cells))
            i = j
        elif s.startswith("## "):
            blocks.append(("heading", [_norm_inline(s[3:].strip())]))
            i += 1
        elif s.startswith("### "):
            blocks.append(("heading", [_norm_inline(s[4:].strip())]))
            i += 1
        elif s.startswith("# "):
            i += 1
        else:
            j = i
            block = []
            while j < len(lines):
                s2 = lines[j].strip()
                if not s2:
                    break
                if s2.startswith(("```", ">", "|", "#")):
                    break
                block.append(lines[j])
                j += 1
            first = block[0].strip() if block else ""
            if (first == "**Q. Geng**\\*"
                    or first.startswith("*Department of Biomedical Engineering")):
                i = j
                continue
            joined = " ".join(ln.strip() for ln in block if ln.strip())
            if re.match(r"^\*\*Figure\s+\d+\.\*\*", joined) and re.search(
                r"\(in\s+`[^`]+\.png`", joined
            ):
                joined = _FIG_ANNOT.sub("", joined)
                joined = re.sub(r"\s{2,}", " ", joined).strip()
                blocks.append(("figure", [_norm_inline(joined)]))
            elif re.match(r"^-\s", joined) or re.match(r"^(\d+)\.\s", joined):
                items = []
                cur = None
                for ln in block:
                    st = ln.strip()
                    if not st:
                        continue
                    if re.match(r"^-\s", st) or re.match(r"^\d+\.\s", st):
                        if cur is not None:
                            items.append(cur)
                        cur = st
                    elif cur is not None:
                        cur += " " + st
                if cur is not None:
                    items.append(cur)
                blocks.append(("list", [_norm_inline(it) for it in items]))
            else:
                blocks.append(("para", [_norm_inline(joined)]))
            i = j
    return blocks


def main() -> int:
    src = SRC.read_text(encoding="utf-8")
    docx_txt = _docx_text()
    blocks = _body_blocks(src)

    missing: list[tuple[str, str]] = []
    n_pieces = 0
    for kind, pieces in blocks:
        for piece in pieces:
            n_pieces += 1
            if piece not in docx_txt:
                missing.append((kind, piece))

    print(f"blocks checked: {len(blocks)}  text pieces: {n_pieces}")
    if missing:
        print(f"MISSING {len(missing)}:")
        for kind, piece in missing:
            print(f"\n[{kind}] {piece[:200]}...")
        return 1
    print("ALL PRESENT — docx text is faithful to manuscript.md")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
